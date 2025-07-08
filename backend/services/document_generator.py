import os
from datetime import datetime
from docxtpl import DocxTemplate, RichText
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import tempfile

class DocumentGenerator:
    def __init__(self):
        self.templates_dir = os.path.join(os.path.dirname(__file__), '..', '..', 'templates')
        os.makedirs(self.templates_dir, exist_ok=True)
        
    def generate(self, agreement, template_name='default'):
        # Create a temporary file for the output
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        
        # Load template if exists, otherwise create from scratch
        template_path = os.path.join(self.templates_dir, f'{template_name}.docx')
        
        if os.path.exists(template_path):
            doc = DocxTemplate(template_path)
            context = self._prepare_context(agreement)
            doc.render(context)
            doc.save(temp_file.name)
        else:
            # Create document from scratch
            doc = Document()
            self._create_document_from_scratch(doc, agreement)
            doc.save(temp_file.name)
        
        return temp_file.name
    
    def _prepare_context(self, agreement):
        """Prepare context dictionary for template rendering"""
        data = agreement.data or {}
        
        # Format dates
        formation_date = agreement.formation_date.strftime('%d %B %Y')
        effective_date = agreement.effective_date.strftime('%d %B %Y')
        
        # Build members table
        members_table = []
        for member in agreement.members:
            members_table.append({
                'name': member.name,
                'entity': member.entity_name or '',
                'class': member.member_class,
                'commitment': f"£{member.capital_commitment:,.0f}" if member.capital_commitment else 'N/A',
                'percentage': f"{member.percentage_interest:.1f}%" if member.percentage_interest else 'TBD'
            })
        
        context = {
            'company_name': agreement.company_name.upper(),
            'state': agreement.state,
            'formation_date': formation_date,
            'effective_date': effective_date,
            'manager_name': agreement.manager_name,
            'manager_entity': agreement.manager_entity or '',
            'principal_place': agreement.principal_place_of_business or '[To be confirmed]',
            'registered_agent': agreement.registered_agent or '[Insert name & address]',
            'purpose': agreement.purpose or data.get('purpose', ''),
            'members_table': members_table,
            **data  # Include all additional data fields
        }
        
        # Handle rich text formatting for certain fields
        if 'title' in context:
            context['title'] = RichText(context['title'], bold=True, size=20)
        
        return context
    
    def _create_document_from_scratch(self, doc, agreement):
        """Create a formatted document from scratch"""
        
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run('LIMITED LIABILITY COMPANY AGREEMENT\nOF\n')
        title_run.bold = True
        title_run.font.size = Pt(16)
        
        company_run = title_para.add_run(f'{agreement.company_name.upper()}\n')
        company_run.bold = True
        company_run.font.size = Pt(18)
        
        subtitle_run = title_para.add_run(f'(a {agreement.state} limited liability company)')
        subtitle_run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Effective Date
        date_para = doc.add_paragraph(f'Effective Date: {agreement.effective_date.strftime("%d %B %Y")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Key Information
        doc.add_paragraph(f'Manager: {agreement.manager_name}' + 
                         (f' ({agreement.manager_entity})' if agreement.manager_entity else ''))
        doc.add_paragraph(f'Principal Place of Business: {agreement.principal_place_of_business or "[To be confirmed]"}')
        doc.add_paragraph(f'Registered Agent ({agreement.state[:2].upper()}): {agreement.registered_agent or "[Insert name & address]"}')
        
        doc.add_paragraph()
        
        # Recitals
        doc.add_heading('Recitals', level=1)
        
        recitals = [
            f'{agreement.company_name} (the "Company") was formed on {agreement.formation_date.strftime("%d %B %Y")} by filing a Certificate of Formation with the {agreement.state} Secretary of State.',
            'The Company has been organized to ' + (agreement.purpose or '[insert purpose]'),
            'The parties desire to enter into this Agreement to govern their rights and obligations as members of the Company.'
        ]
        
        for i, recital in enumerate(recitals):
            doc.add_paragraph(f'{chr(65+i)}. {recital}')
        
        doc.add_paragraph()
        doc.add_paragraph('NOW, THEREFORE, in consideration of the mutual covenants herein, the parties agree as follows:')
        
        # Articles
        self._add_article_1(doc, agreement)
        self._add_article_2(doc, agreement)
        self._add_article_3(doc, agreement)
        self._add_article_4(doc)
        self._add_article_5(doc, agreement)
        self._add_article_6(doc, agreement)
        self._add_signature_page(doc, agreement)
    
    def _add_article_1(self, doc, agreement):
        doc.add_heading('ARTICLE I – Definitions & Construction', level=1)
        doc.add_paragraph('Key defined terms are set out in Schedule A. Where terms are not defined, they have the meaning given in the ' +
                         f'{agreement.state} Limited Liability Company Act.')
        doc.add_paragraph()
    
    def _add_article_2(self, doc, agreement):
        doc.add_heading('ARTICLE II – Formation, Purpose, Term', level=1)
        
        doc.add_paragraph('2.1 Formation & Name.', style='Heading 2')
        doc.add_paragraph(f'The Company exists as a {agreement.state} LLC under the name {agreement.company_name} ' +
                         'and may operate under trade or "doing-business-as" names approved by the Manager.')
        
        doc.add_paragraph('2.2 Purpose.', style='Heading 2')
        doc.add_paragraph("The Company's purpose is limited to: " + (agreement.purpose or '[insert purpose]'))
        
        doc.add_paragraph('2.3 Term.', style='Heading 2')
        doc.add_paragraph('Perpetual, unless dissolved under Article X.')
        doc.add_paragraph()
    
    def _add_article_3(self, doc, agreement):
        doc.add_heading('ARTICLE III – Units, Capitalization & Classes', level=1)
        
        doc.add_paragraph('3.1 Authorized Units.', style='Heading 2')
        doc.add_paragraph('The Company is authorized to issue three classes of limited liability company interests (collectively, "Units"):')
        
        # Create units table
        table = doc.add_table(rows=4, cols=4)
        table.style = 'Light List Accent 1'
        
        # Header row
        headers = ['Class', 'Designation', 'Pre-Money Valuation', 'Rights Snapshot']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        # Data rows - from agreement data
        data = agreement.data or {}
        class_data = [
            ['Class A', 'Anchor Units', data.get('class_a_valuation', '£0.95m'), 
             data.get('class_a_rights', '≥1 Board seat; veto over Reserved Matters')],
            ['Class B', 'Investor Units', data.get('class_b_valuation', '£3.25m'), 
             data.get('class_b_rights', 'Standard voting; pro-rata pre-emptive rights')],
            ['Class C', 'Sweat-Equity Units', 'N/A (services)', 
             data.get('class_c_rights', 'Vesting schedules tied to KPIs')]
        ]
        
        for i, row_data in enumerate(class_data):
            for j, cell_text in enumerate(row_data):
                table.rows[i+1].cells[j].text = cell_text
        
        doc.add_paragraph()
        
        doc.add_paragraph('3.2 Initial Capital Commitments.', style='Heading 2')
        
        # Members table
        if agreement.members:
            members_table = doc.add_table(rows=len(agreement.members)+1, cols=4)
            members_table.style = 'Light List Accent 1'
            
            # Header
            headers = ['Member', 'Class', 'Commitment', 'Payment Terms']
            for i, header in enumerate(headers):
                cell = members_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            # Member rows
            for i, member in enumerate(agreement.members):
                members_table.rows[i+1].cells[0].text = member.name
                members_table.rows[i+1].cells[1].text = member.member_class
                members_table.rows[i+1].cells[2].text = f"£{member.capital_commitment:,.0f}" if member.capital_commitment else 'Services'
                members_table.rows[i+1].cells[3].text = 'See Schedule C' if member.member_class == 'A' else 'Pro-rata calls'
        
        doc.add_paragraph()
    
    def _add_article_4(self, doc):
        doc.add_heading('ARTICLE IV – Allocations & Tax', level=1)
        doc.add_paragraph('Standard tax provisions apply, with profits and losses allocated pro-rata to Percentage Interests.')
        doc.add_paragraph()
    
    def _add_article_5(self, doc, agreement):
        doc.add_heading('ARTICLE V – Distributions & Waterfall', level=1)
        
        doc.add_paragraph('5.1 Timing.', style='Heading 2')
        doc.add_paragraph('Distributions are at Manager discretion, subject to lender covenants and cash-flow needs.')
        
        doc.add_paragraph('5.2 Waterfall.', style='Heading 2')
        doc.add_paragraph('Distributable cash (including exit proceeds) is applied:')
        
        waterfall_items = [
            'Transaction Costs & Liabilities.',
            'Return of Capital. Repay Members pro-rata until all Capital Contributions returned.',
            'Catch-Up / Carry. 20% to Manager until it has received 20% of total distributed amounts above return of capital.',
            'Residual. 80% to all Members pro-rata by fully-diluted ownership; 20% to Manager (carry).'
        ]
        
        for i, item in enumerate(waterfall_items, 1):
            doc.add_paragraph(f'{i}. {item}', style='List Number')
        
        doc.add_paragraph()
    
    def _add_article_6(self, doc, agreement):
        doc.add_heading('ARTICLE VI – Governance', level=1)
        
        doc.add_paragraph('6.1 Manager Powers.', style='Heading 2')
        doc.add_paragraph(f'The Manager ({agreement.manager_name}) has exclusive authority over operations, ' +
                         'subject only to Reserved Matters.')
        
        doc.add_paragraph('6.2 Reserved Matters.', style='Heading 2')
        doc.add_paragraph('No action without requisite Member consent on:')
        
        reserved_matters = [
            'Issuance of equity or options outside approved pools.',
            'Incurrence of new secured debt above specified thresholds.',
            'Sale or encumbrance of material assets.',
            'Key personnel decisions.'
        ]
        
        for matter in reserved_matters:
            doc.add_paragraph(f'• {matter}', style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_signature_page(self, doc, agreement):
        doc.add_page_break()
        doc.add_heading('Signature Page', level=1)
        
        doc.add_paragraph('By signing below, each undersigned Person agrees to be bound by this ' +
                         f'Limited Liability Company Agreement of {agreement.company_name} as of the Effective Date.')
        doc.add_paragraph()
        
        # Signature table
        sig_table = doc.add_table(rows=len(agreement.members)+2, cols=3)
        sig_table.style = 'Table Grid'
        
        # Header
        headers = ['Member / Manager', 'Capacity', 'Signature / Date']
        for i, header in enumerate(headers):
            cell = sig_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        # Manager row
        sig_table.rows[1].cells[0].text = agreement.manager_name
        sig_table.rows[1].cells[1].text = 'Manager'
        sig_table.rows[1].cells[2].text = '_' * 30
        
        # Member rows
        for i, member in enumerate(agreement.members):
            sig_table.rows[i+2].cells[0].text = member.name
            sig_table.rows[i+2].cells[1].text = f'Member (Class {member.member_class})'
            sig_table.rows[i+2].cells[2].text = '_' * 30